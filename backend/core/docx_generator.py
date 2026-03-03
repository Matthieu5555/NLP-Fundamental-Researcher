"""
DOCX Report Generator — produces editable Word documents.

Uses python-docx to build a professional equity research report
programmatically. Every element (paragraphs, tables, images) is a native
Word object, fully editable in Microsoft Word or Google Docs.

Single entry point: generate_docx() -> bytes
"""

import io
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, TYPE_CHECKING

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from .report_builder import ReportState, Source, AnalystSource
from .branding_config import ReportBrandingConfig, get_default_config
from .pdf_formatting import (
    extract_highlights,
    extract_recommendation,
    format_number,
    format_percent,
    get_exchange_name,
    strip_markdown_tables,
    strip_markup,
)
from .valuation_display import build_valuation_display

if TYPE_CHECKING:
    from src_george_researcher.data_fetchers.stock_data import StockInfo
    from .pdf_data_collector import FirstPageData

logger = logging.getLogger(__name__)

# Brand blue — sourced from branding config at runtime, this is the fallback
_DEFAULT_BRAND_COLOR = RGBColor(0x1E, 0x3A, 0x5F)

# Typography constants
_FONT_BODY = "Calibri"
_FONT_SIZE_BODY = Pt(10)
_FONT_SIZE_SMALL = Pt(8)
_FONT_SIZE_H1 = Pt(20)
_FONT_SIZE_H2 = Pt(14)
_FONT_SIZE_H3 = Pt(11)
_FONT_SIZE_METRIC_LABEL = Pt(8)
_FONT_SIZE_METRIC_VALUE = Pt(10)

_COLOR_BODY = RGBColor(0x33, 0x33, 0x33)
_COLOR_SECONDARY = RGBColor(0x66, 0x66, 0x66)
_COLOR_POSITIVE = RGBColor(0x05, 0x96, 0x69)
_COLOR_NEGATIVE = RGBColor(0xDC, 0x26, 0x26)
_COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_COLOR_LIGHT_BG = RGBColor(0xF8, 0xFA, 0xFC)


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert '#1E3A5F' to RGBColor."""
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color.lstrip("#")}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_shading_para(paragraph, hex_color: str) -> None:
    """Apply background shading to a paragraph (for inline shading within cells)."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color.lstrip("#")}" w:val="clear"/>'
    )
    paragraph._p.get_or_add_pPr().append(shading)


def _add_sidebar_section(cell, title: str, rows: List[tuple],
                         brand_color: RGBColor, hex_brand: str) -> None:
    """Add a compact sidebar data section (title + label/value rows) into a cell.

    Mimics the PDF's .sidebar-section: uppercase header with bottom border,
    followed by dotted-separator data rows with label left, value right.
    """
    # Section header
    hp = cell.add_paragraph()
    _add_run(hp, title.upper(), bold=True, size=Pt(8), color=brand_color)
    _set_paragraph_spacing(hp, before=Pt(8), after=Pt(3))
    hp._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{hex_brand}"/>'
        '</w:pBdr>'
    ))

    # Data rows as tab-separated label/value pairs
    for label, value in rows:
        rp = cell.add_paragraph()
        rp.paragraph_format.tab_stops.add_tab_stop(Cm(4.5), WD_TAB_ALIGNMENT.RIGHT)
        _add_run(rp, label, size=Pt(8), color=_COLOR_BODY)
        _add_run(rp, "\t", size=Pt(8))
        _add_run(rp, str(value), bold=True, size=Pt(8), color=_COLOR_BODY)
        _set_paragraph_spacing(rp, before=Pt(1), after=Pt(1))
        # Dotted bottom border between rows
        rp._p.get_or_add_pPr().append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '<w:bottom w:val="dotted" w:sz="2" w:space="1" w:color="CCCCCC"/>'
            '</w:pBdr>'
        ))


def _add_accent_box(doc: Document, bg_hex: str, border_hex: str,
                    border_size: int = 24) -> 'Cell':
    """Create an accent box (colored left border + background) and return the cell.

    Builds a single-cell borderless table with background shading and a thick
    colored left border.  The caller writes content into the returned Cell via
    cell.add_paragraph().  All OOXML plumbing is encapsulated here.

    Args:
        doc: The Document to append the table to.
        bg_hex: Background color, e.g. '#dcfce7'.
        border_hex: Left border color, e.g. '#16a34a'.
        border_size: Border width in half-points (24 = 3pt).

    Returns:
        The table Cell ready for content insertion.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl

    # Remove all borders, then apply only the left accent border on the cell
    no_borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>'
    )
    tbl_pr.append(no_borders)

    cell = table.cell(0, 0)
    _set_cell_shading(cell, bg_hex)

    # Thick colored left border on the cell
    tc_pr = cell._tc.get_or_add_tcPr()
    border_clean = border_hex.lstrip("#")
    tc_pr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_clean}"/>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    ))

    # Cell padding for breathing room
    tc_pr.append(parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        '<w:top w:w="100" w:type="dxa"/>'
        '<w:left w:w="200" w:type="dxa"/>'
        '<w:bottom w:w="100" w:type="dxa"/>'
        '<w:right w:w="160" w:type="dxa"/>'
        '</w:tcMar>'
    ))

    # Remove the default empty paragraph so caller starts clean
    for p in cell.paragraphs:
        p.clear()

    return cell


def _add_run(paragraph, text: str, bold: bool = False, italic: bool = False,
             size: Pt = None, color: RGBColor = None, font_name: str = None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name or _FONT_BODY
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def _set_paragraph_spacing(paragraph, before: Pt = None, after: Pt = None,
                           line_spacing: float = None):
    """Set paragraph spacing."""
    fmt = paragraph.paragraph_format
    if before is not None:
        fmt.space_before = before
    if after is not None:
        fmt.space_after = after
    if line_spacing is not None:
        fmt.line_spacing = line_spacing


def _add_page_number_run(paragraph, size: Pt = Pt(7),
                         color: RGBColor = None) -> None:
    """Insert a PAGE field code into the paragraph (renders as page number in Word)."""
    run = paragraph.add_run()
    run.font.size = size
    run.font.name = _FONT_BODY
    if color:
        run.font.color.rgb = color

    # Word PAGE field: fldChar begin, instrText, fldChar end
    from docx.oxml import OxmlElement
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'begin')
    run._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' PAGE '
    run._r.append(instr)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'end')
    run._r.append(fld_end)


def _add_cover_header_bar(
    doc: Document,
    company_name: str,
    ticker: str,
    exchange: str,
    sector: str,
    rating: str,
    price_target: Optional[float],
    target_upside: Optional[float],
    current_price: Optional[float],
    brand_color: RGBColor,
    hex_brand: str,
) -> None:
    """Render the full-width branded header bar on the cover page.

    Creates a two-column borderless table shaded in the brand color.
    Left cell: company name, ticker, exchange, sector in white.
    Right cell: rating badge, price target, upside/downside.
    """
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl

    # Remove all table borders
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>'
    )
    tbl_pr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    ))

    padding_xml = (
        f'<w:tcMar {nsdecls("w")}>'
        '<w:top w:w="140" w:type="dxa"/>'
        '<w:left w:w="220" w:type="dxa"/>'
        '<w:bottom w:w="140" w:type="dxa"/>'
        '<w:right w:w="220" w:type="dxa"/>'
        '</w:tcMar>'
    )

    # --- Left cell: company info ---
    left = table.cell(0, 0)
    _set_cell_shading(left, hex_brand)
    left._tc.get_or_add_tcPr().append(parse_xml(padding_xml))

    # Clear default paragraph, then write content
    lp = left.paragraphs[0]
    lp.clear()
    _add_run(lp, company_name, bold=True, size=Pt(22), color=_COLOR_WHITE)
    _set_paragraph_spacing(lp, after=Pt(2))

    lp2 = left.add_paragraph()
    parts = [ticker]
    if exchange:
        parts.append(get_exchange_name(exchange))
    if sector:
        parts.append(sector)
    _add_run(lp2, "  |  ".join(parts), size=Pt(10), color=_COLOR_WHITE)
    _set_paragraph_spacing(lp2, before=Pt(0), after=Pt(0))

    # --- Right cell: rating + price target ---
    right = table.cell(0, 1)
    _set_cell_shading(right, hex_brand)
    right._tc.get_or_add_tcPr().append(parse_xml(padding_xml))

    rp = right.paragraphs[0]
    rp.clear()
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(rp, rating, bold=True, size=Pt(16), color=_COLOR_WHITE)
    _set_paragraph_spacing(rp, after=Pt(2))

    if price_target or current_price:
        rp2 = right.add_paragraph()
        rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if price_target:
            _add_run(rp2, f"Target ${price_target:,.2f}", bold=True, size=Pt(11),
                     color=_COLOR_WHITE)
            if target_upside is not None:
                _add_run(rp2, f"  ({target_upside:+.1f}%)", size=Pt(10),
                         color=_COLOR_WHITE)
        if current_price:
            _add_run(rp2, f"   Current ${current_price:,.2f}", size=Pt(9),
                     color=_COLOR_WHITE)
        _set_paragraph_spacing(rp2, before=Pt(0), after=Pt(0))

    # Set approximate column widths (70/30 split of ~17cm usable width)
    from docx.shared import Emu
    for cell_obj, width_cm in [(left, 12.0), (right, 5.0)]:
        cell_obj.width = Cm(width_cm)


# ---------------------------------------------------------------------------
# Markdown-to-DOCX conversion
# ---------------------------------------------------------------------------

def _strip_markdown_images(text: str) -> str:
    """Remove markdown image references like ![alt](url)."""
    return re.sub(r'!\[.*?\]\(.*?\)', '', text)


def _write_markdown_to_doc(doc: Document, md_text: str, brand_color: RGBColor) -> None:
    """
    Parse markdown text and write it as native Word paragraphs.

    Handles: headings (##, ###), bold (**), italic (*), bullet lists (-),
    numbered lists, and plain paragraphs. Tables in markdown are converted
    to Word tables.
    """
    if not md_text or not md_text.strip():
        return

    md_text = _strip_markdown_images(md_text)
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Markdown table detection
        if "|" in line and i + 1 < len(lines) and re.match(r'\s*\|[\s\-:|]+\|', lines[i + 1]):
            table_lines = [line]
            i += 1
            while i < len(lines) and "|" in lines[i]:
                if not re.match(r'\s*\|[\s\-:|]+\|', lines[i]):
                    table_lines.append(lines[i])
                i += 1
            _write_markdown_table(doc, table_lines, brand_color)
            continue

        # Headings — use proper Word heading styles for outline and ToC
        if line.startswith("#### "):
            doc.add_paragraph(line[5:].strip(), style='Heading 3')
            i += 1
            continue

        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style='Heading 3')
            i += 1
            continue

        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style='Heading 2')
            i += 1
            continue

        if line.startswith("# "):
            doc.add_paragraph(line[2:].strip(), style='Heading 2')
            i += 1
            continue

        # Bullet lists
        if re.match(r'^\s*[-*]\s', line):
            p = doc.add_paragraph(style="List Bullet")
            _write_inline_formatting(p, re.sub(r'^\s*[-*]\s+', '', line))
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\s*\d+\.\s', line):
            p = doc.add_paragraph(style="List Number")
            _write_inline_formatting(p, re.sub(r'^\s*\d+\.\s+', '', line))
            i += 1
            continue

        # Regular paragraph — collect continuation lines
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") \
                and not re.match(r'^\s*[-*]\s', lines[i]) \
                and not re.match(r'^\s*\d+\.\s', lines[i]) \
                and "|" not in lines[i]:
            para_lines.append(lines[i])
            i += 1

        p = doc.add_paragraph()
        _write_inline_formatting(p, " ".join(para_lines))
        _set_paragraph_spacing(p, after=Pt(6))


def _write_inline_formatting(paragraph, text: str) -> None:
    """Apply bold and italic inline formatting from markdown syntax."""
    # Process **bold** and *italic* markers
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)')
    last_end = 0

    for match in pattern.finditer(text):
        # Add text before this match
        if match.start() > last_end:
            _add_run(paragraph, text[last_end:match.start()], size=_FONT_SIZE_BODY, color=_COLOR_BODY)

        if match.group(2):  # ***bold italic***
            _add_run(paragraph, match.group(2), bold=True, italic=True,
                     size=_FONT_SIZE_BODY, color=_COLOR_BODY)
        elif match.group(3):  # **bold**
            _add_run(paragraph, match.group(3), bold=True, size=_FONT_SIZE_BODY, color=_COLOR_BODY)
        elif match.group(4):  # *italic*
            _add_run(paragraph, match.group(4), italic=True, size=_FONT_SIZE_BODY, color=_COLOR_BODY)

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        _add_run(paragraph, text[last_end:], size=_FONT_SIZE_BODY, color=_COLOR_BODY)


def _write_markdown_table(doc: Document, table_lines: List[str],
                          brand_color: RGBColor) -> None:
    """Convert markdown table lines into a Word table."""
    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows_data.append(cells)

    if len(rows_data) < 2:
        return

    headers = rows_data[0]
    data_rows = rows_data[1:]
    n_cols = len(headers)

    table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        _add_run(p, header, bold=True, size=_FONT_SIZE_SMALL, color=_COLOR_WHITE)
        _set_cell_shading(cell, brand_color.__str__().replace("(", "").replace(")", "").replace(", ", ""))
        # Use hex directly
        hex_color = f"{brand_color[0]:02x}{brand_color[1]:02x}{brand_color[2]:02x}"
        _set_cell_shading(cell, hex_color)

    # Data rows
    for i, row in enumerate(data_rows):
        for j, val in enumerate(row[:n_cols]):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_run(p, val, size=_FONT_SIZE_SMALL, color=_COLOR_BODY)
            # Alternating row shading
            if i % 2 == 0:
                _set_cell_shading(cell, "F8FAFC")


# ---------------------------------------------------------------------------
# Structured valuation tables
# ---------------------------------------------------------------------------

def _add_section_heading(doc: Document, title: str, brand_color: RGBColor,
                         page_break: bool = False) -> None:
    """Add a styled section heading using Heading 1 style with bottom border accent."""
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph(title, style='Heading 1')

    # Bottom border line in brand color
    hex_clean = f"{brand_color[0]:02x}{brand_color[1]:02x}{brand_color[2]:02x}"
    p._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="{hex_clean}"/>'
        '</w:pBdr>'
    ))


def _add_key_value_table(doc: Document, items: List[Dict[str, str]],
                         brand_color: RGBColor, title: str = None) -> None:
    """Add a two-column key-value table (label | value)."""
    if not items:
        return

    if title:
        p = doc.add_paragraph()
        _add_run(p, title, bold=True, size=_FONT_SIZE_H3, color=brand_color)
        _set_paragraph_spacing(p, before=Pt(8), after=Pt(4))

    table = doc.add_table(rows=len(items), cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(2.0)

    for i, item in enumerate(items):
        label_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]

        label_cell.text = ""
        value_cell.text = ""

        _add_run(label_cell.paragraphs[0], str(item.get("label", "")),
                 size=_FONT_SIZE_SMALL, color=_COLOR_SECONDARY)
        _add_run(value_cell.paragraphs[0], str(item.get("value", "N/A")),
                 bold=True, size=_FONT_SIZE_METRIC_VALUE, color=_COLOR_BODY)
        value_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if i % 2 == 0:
            _set_cell_shading(label_cell, "F8FAFC")
            _set_cell_shading(value_cell, "F8FAFC")


def _add_scenario_table(doc: Document, scenario_summary: Dict,
                        brand_color: RGBColor) -> None:
    """Add scenario analysis as a Word table."""
    if not scenario_summary:
        return

    headers = ["Scenario", "Fair Value", "Upside/Downside", "Probability"]
    rows = []
    for label in ("bull", "base", "bear"):
        s = scenario_summary.get(label)
        if not s:
            continue
        fv = s.get("fair_value")
        upside = s.get("upside_pct")
        prob = s.get("probability")
        rows.append([
            label.capitalize(),
            f"${fv:,.2f}" if fv else "N/A",
            f"{upside:+.1f}%" if upside is not None else "N/A",
            f"{prob:.0%}" if prob else "N/A",
        ])

    weighted = scenario_summary.get("weighted_fair_value")
    if weighted:
        rows.append(["Weighted Value", f"${weighted:,.2f}", "", ""])

    _add_data_table(doc, headers, rows, brand_color)


def _add_sensitivity_grid_table(doc: Document, grid_data: Dict,
                                brand_color: RGBColor,
                                title: str = "Sensitivity Analysis") -> None:
    """Add a sensitivity analysis grid as a Word table."""
    if not grid_data:
        return

    row_vals = grid_data.get("row_values", [])
    col_vals = grid_data.get("col_values", [])
    grid = grid_data.get("grid", [])
    cell_classes = grid_data.get("cell_classes", [])

    if not row_vals or not col_vals or not grid:
        return

    p = doc.add_paragraph()
    _add_run(p, title, bold=True, size=_FONT_SIZE_H3, color=brand_color)
    _set_paragraph_spacing(p, before=Pt(8), after=Pt(4))

    n_cols = len(col_vals) + 1
    n_rows = len(row_vals) + 1

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row: empty corner + column values
    corner = table.rows[0].cells[0]
    corner.text = ""
    hex_brand = f"{brand_color[0]:02x}{brand_color[1]:02x}{brand_color[2]:02x}"
    _set_cell_shading(corner, hex_brand)

    for j, cv in enumerate(col_vals):
        cell = table.rows[0].cells[j + 1]
        cell.text = ""
        cv_display = format_percent(cv) if isinstance(cv, (int, float)) else f"{cv}"
        _add_run(cell.paragraphs[0], cv_display, bold=True, size=_FONT_SIZE_SMALL, color=_COLOR_WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, hex_brand)

    # Data rows
    for i, rv in enumerate(row_vals):
        # Row header
        row_header = table.rows[i + 1].cells[0]
        row_header.text = ""
        rv_display = format_percent(rv) if isinstance(rv, (int, float)) else f"{rv}"
        _add_run(row_header.paragraphs[0], rv_display, bold=True, size=_FONT_SIZE_SMALL, color=brand_color)

        for j, val in enumerate(grid[i] if i < len(grid) else []):
            cell = table.rows[i + 1].cells[j + 1]
            cell.text = ""

            display_val = f"${val:,.2f}" if isinstance(val, (int, float)) else str(val)
            cls = cell_classes[i][j] if i < len(cell_classes) and j < len(cell_classes[i]) else ""

            color = _COLOR_BODY
            if cls == "positive":
                color = _COLOR_POSITIVE
            elif cls == "negative":
                color = _COLOR_NEGATIVE
            elif cls == "base":
                color = brand_color
                _set_cell_shading(cell, "EFF6FF")

            _add_run(cell.paragraphs[0], display_val, size=_FONT_SIZE_SMALL, color=color,
                     bold=(cls == "base"))
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_football_field_table(doc: Document, ranges: List[Dict],
                              current_price: Optional[float],
                              brand_color: RGBColor) -> None:
    """Add football field valuation ranges as a table."""
    if not ranges:
        return

    headers = ["Method", "Low", "Mid", "High"]
    rows = []
    for r in ranges:
        rows.append([
            r.get("method", ""),
            f"${r['low']:,.2f}" if r.get("low") else "N/A",
            f"${r['mid']:,.2f}" if r.get("mid") else "N/A",
            f"${r['high']:,.2f}" if r.get("high") else "N/A",
        ])

    if current_price:
        rows.append(["Current Price", f"${current_price:,.2f}", "", ""])

    _add_data_table(doc, headers, rows, brand_color)


def _add_comps_table(doc: Document, comp_summary: Dict,
                     brand_color: RGBColor) -> None:
    """Add comparable companies table."""
    if not comp_summary:
        return

    peers = comp_summary.get("peers", [])
    target = comp_summary.get("target", {})
    medians = comp_summary.get("medians", {})

    if not peers:
        return

    # Determine available metrics
    metric_keys = []
    metric_labels = {
        "pe_ratio": "P/E",
        "ev_ebitda": "EV/EBITDA",
        "ev_revenue": "EV/Revenue",
        "price_to_book": "P/B",
        "roe": "ROE",
    }
    for key in metric_keys:
        if any(p.get(key) is not None for p in peers):
            metric_keys.append(key)

    if not metric_keys:
        # Default common metrics
        metric_keys = ["pe_ratio", "ev_ebitda", "ev_revenue"]

    headers = ["Company"] + [metric_labels.get(k, k) for k in metric_keys]
    rows = []

    # Target company first
    if target:
        row = [target.get("name", target.get("ticker", "Target"))]
        for k in metric_keys:
            v = target.get(k)
            row.append(f"{v:.1f}x" if isinstance(v, (int, float)) else "N/A")
        rows.append(row)

    # Peers
    for peer in peers:
        row = [peer.get("name", peer.get("ticker", ""))]
        for k in metric_keys:
            v = peer.get(k)
            row.append(f"{v:.1f}x" if isinstance(v, (int, float)) else "N/A")
        rows.append(row)

    # Medians
    if medians:
        row = ["Median"]
        for k in metric_keys:
            v = medians.get(k)
            row.append(f"{v:.1f}x" if isinstance(v, (int, float)) else "N/A")
        rows.append(row)

    _add_data_table(doc, headers, rows, brand_color)


def _add_earnings_table(doc: Document, earnings_data: Dict,
                        brand_color: RGBColor) -> None:
    """Add earnings model table, filtering out rows where all value fields are null."""
    if not earnings_data:
        return

    rows_data = earnings_data.get("rows", [])
    if not rows_data:
        return

    # Filter out rows where all numeric fields are None (e.g. fiscal_year: 0 placeholders)
    _value_fields = ("revenue", "ebitda", "net_income", "eps", "revenue_growth", "ebitda_margin")
    rows_data = [
        r for r in rows_data
        if any(r.get(k) is not None for k in _value_fields)
    ]
    if not rows_data:
        return

    # Determine columns from first row
    sample = rows_data[0]
    headers = ["Period"]
    value_keys = []
    for key in ("revenue", "ebitda", "net_income", "eps", "revenue_growth", "ebitda_margin"):
        if key in sample:
            label_map = {
                "revenue": "Revenue",
                "ebitda": "EBITDA",
                "net_income": "Net Income",
                "eps": "EPS",
                "revenue_growth": "Rev Growth",
                "ebitda_margin": "EBITDA Margin",
            }
            headers.append(label_map.get(key, key))
            value_keys.append(key)

    rows = []
    for r in rows_data:
        row = [str(r.get("period", r.get("year", "")))]
        is_estimate = r.get("is_estimate", False)
        for k in value_keys:
            v = r.get(k)
            if v is None:
                row.append("N/A")
            elif k in ("revenue_growth", "ebitda_margin"):
                row.append(format_percent(v))
            elif k == "eps":
                row.append(f"${v:.2f}")
            elif isinstance(v, (int, float)):
                row.append(f"${v:,.0f}" if abs(v) > 100 else f"${v:.2f}")
            else:
                row.append(str(v))
        if is_estimate:
            row[0] = row[0] + "E"
        rows.append(row)

    _add_data_table(doc, headers, rows, brand_color)


def _add_dcf_summary_table(doc: Document, dcf_summary: Dict,
                           brand_color: RGBColor) -> None:
    """Add DCF assumptions and output as tables."""
    if not dcf_summary:
        return

    # Key outputs
    items = []
    fv = dcf_summary.get("fair_value")
    if fv:
        items.append({"label": "Fair Value per Share", "value": f"${fv:,.2f}"})
    upside = dcf_summary.get("upside_pct")
    if upside is not None:
        items.append({"label": "Upside/Downside", "value": f"{upside:+.1f}%"})

    wacc = dcf_summary.get("wacc")
    if wacc:
        items.append({"label": "WACC", "value": format_percent(wacc)})

    tgr = dcf_summary.get("terminal_growth_rate")
    if tgr:
        items.append({"label": "Terminal Growth Rate", "value": format_percent(tgr)})

    ev = dcf_summary.get("enterprise_value")
    if ev:
        items.append({"label": "Enterprise Value", "value": f"${ev:,.0f}"})

    tv_pct = dcf_summary.get("tv_pct_of_ev")
    if tv_pct:
        items.append({"label": "Terminal Value % of EV", "value": f"{tv_pct:.1f}%"})

    if items:
        _add_key_value_table(doc, items, brand_color, title="DCF Key Outputs")

    # Projection details
    projections = dcf_summary.get("projection_details", [])
    if projections:
        headers = ["Year"]
        sample = projections[0]
        value_keys = []
        for k in ("revenue", "fcf", "discounted_fcf"):
            if k in sample:
                label_map = {"revenue": "Revenue", "fcf": "FCF", "discounted_fcf": "Discounted FCF"}
                headers.append(label_map.get(k, k))
                value_keys.append(k)

        rows = []
        for proj in projections:
            row = [str(proj.get("year", ""))]
            for k in value_keys:
                v = proj.get(k)
                row.append(f"${v:,.0f}" if isinstance(v, (int, float)) else "N/A")
            rows.append(row)

        if rows:
            p = doc.add_paragraph()
            _add_run(p, "Projected Cash Flows", bold=True, size=_FONT_SIZE_H3, color=brand_color)
            _set_paragraph_spacing(p, before=Pt(8), after=Pt(4))
            _add_data_table(doc, headers, rows, brand_color)


def _add_data_table(doc: Document, headers: List[str], rows: List[List[str]],
                    brand_color: RGBColor) -> None:
    """Add a generic data table with branded header row and alternating shading."""
    if not rows:
        return

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hex_brand = f"{brand_color[0]:02x}{brand_color[1]:02x}{brand_color[2]:02x}"

    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        _add_run(cell.paragraphs[0], header, bold=True, size=_FONT_SIZE_SMALL, color=_COLOR_WHITE)
        _set_cell_shading(cell, hex_brand)
        if j > 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row[:n_cols]):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            _add_run(cell.paragraphs[0], strip_markup(str(val)),
                     size=_FONT_SIZE_SMALL, color=_COLOR_BODY)
            if j > 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if i % 2 == 0:
                _set_cell_shading(cell, "F8FAFC")

    doc.add_paragraph()  # spacing after table


def _add_financial_statements(doc: Document, financial_statements: Dict,
                              brand_color: RGBColor) -> None:
    """Add financial statements (IS, BS, CF) as Word tables."""
    if not financial_statements:
        return

    statement_names = {
        "income_statement": "Income Statement",
        "balance_sheet": "Balance Sheet",
        "cash_flow": "Cash Flow Statement",
    }

    for key, title in statement_names.items():
        stmt = financial_statements.get(key)
        if not stmt:
            continue

        p = doc.add_paragraph()
        _add_run(p, title, bold=True, size=_FONT_SIZE_H3, color=brand_color)
        _set_paragraph_spacing(p, before=Pt(10), after=Pt(4))

        # Financial statements are typically {metric: {year: value}}
        if isinstance(stmt, dict):
            # Get years from first metric
            years = []
            for metric_data in stmt.values():
                if isinstance(metric_data, dict):
                    years = sorted(metric_data.keys(), reverse=True)[:5]
                    break

            if years:
                headers = ["Metric"] + [str(y) for y in years]
                rows = []
                for metric, year_data in stmt.items():
                    if isinstance(year_data, dict):
                        row = [str(metric)]
                        for y in years:
                            v = year_data.get(y)
                            if isinstance(v, (int, float)):
                                row.append(f"{v:,.0f}" if abs(v) > 100 else f"{v:.2f}")
                            else:
                                row.append(str(v) if v is not None else "N/A")
                        rows.append(row)

                if rows:
                    _add_data_table(doc, headers, rows, brand_color)


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def generate_docx(
    report_state: ReportState,
    ticker: str,
    stock_info: Optional["StockInfo"] = None,
    branding: Optional[ReportBrandingConfig] = None,
    analyst_sources: Optional[List[AnalystSource]] = None,
    financial_statements: Optional[Dict[str, Any]] = None,
    session_metadata: Optional[Dict[str, Any]] = None,
) -> bytes:
    """
    Generate a fully editable Word document from the analysis.

    Produces a professional equity research report with native Word tables,
    paragraphs, and embedded chart images. Every element is directly editable
    in Microsoft Word or Google Docs.

    Args:
        report_state: The report data with sections and metadata.
        ticker: Stock ticker symbol.
        stock_info: Stock information (price, market cap, etc.).
        branding: Branding configuration (firm, analyst).
        analyst_sources: List of analyst belief sources.
        financial_statements: Financial statement data.
        session_metadata: Full session metadata with valuation outputs.

    Returns:
        DOCX file as bytes.
    """
    branding = branding or get_default_config()
    brand_color = _hex_to_rgb(branding.firm.primary_color)
    hex_brand = branding.firm.primary_color.lstrip("#")

    doc = Document()

    # Configure default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = _FONT_BODY
    font.size = _FONT_SIZE_BODY
    font.color.rgb = _COLOR_BODY

    # Configure heading styles for proper document outline and ToC support
    for level, size, sp_before, sp_after in [
        (1, _FONT_SIZE_H1, Pt(18), Pt(8)),
        (2, _FONT_SIZE_H2, Pt(14), Pt(6)),
        (3, _FONT_SIZE_H3, Pt(10), Pt(4)),
    ]:
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = _FONT_BODY
        h_style.font.size = size
        h_style.font.bold = True
        h_style.font.color.rgb = brand_color
        h_style.paragraph_format.space_before = sp_before
        h_style.paragraph_format.space_after = sp_after
        h_style.paragraph_format.keep_with_next = True

    # Page margins and footer
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        # Running footer: firm name | page number | disclaimer
        section.different_first_page_header_footer = True
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.paragraph_format.tab_stops.add_tab_stop(Cm(8.5), WD_TAB_ALIGNMENT.CENTER)
        fp.paragraph_format.tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)
        _add_run(fp, branding.firm.name, size=Pt(7), color=_COLOR_SECONDARY)
        _add_run(fp, "\t", size=Pt(7))
        # Page number via field code
        _add_page_number_run(fp, size=Pt(7), color=_COLOR_SECONDARY)
        _add_run(fp, "\tSee disclosures at end of report", size=Pt(7),
                 color=_COLOR_SECONDARY)

    # ===== Collect data (mirroring PDF pipeline) =====

    first_page_data = None
    try:
        from .pdf_data_collector import collect_first_page_data
        first_page_data = collect_first_page_data(ticker, stock_info)
    except Exception as e:
        logger.warning("Could not collect first page data: %s", e)

    rating = extract_recommendation(report_state)
    current_price = None

    # Company info
    company_name = ticker
    exchange = ""
    sector = ""

    if stock_info:
        company_name = getattr(stock_info, "name", ticker) or ticker
        exchange = getattr(stock_info, "exchange", "") or ""
        sector = getattr(stock_info, "sector", "") or ""
    elif first_page_data:
        company_name = first_page_data.company_name or ticker
        exchange = first_page_data.exchange or ""
        sector = first_page_data.sector or ""

    if first_page_data:
        current_price = first_page_data.current_price

    # Valuation display data
    from .valuation_display import ValuationDisplay

    val_display = ValuationDisplay()
    if session_metadata:
        val_display = build_valuation_display(session_metadata, current_price=current_price)

    # Price target
    price_target = None
    target_upside = None
    scenario_sum = val_display.scenario_summary
    if scenario_sum.get("weighted_fair_value"):
        price_target = scenario_sum["weighted_fair_value"]
    elif val_display.dcf_summary.get("fair_value"):
        price_target = val_display.dcf_summary["fair_value"]

    if price_target and current_price and current_price > 0:
        target_upside = round(((price_target - current_price) / current_price) * 100, 1)

    # Report headline
    highlights_raw = []
    thesis_md = ""
    if "investment_thesis" in report_state.sections:
        sec = report_state.sections["investment_thesis"]
        if hasattr(sec, "content") and sec.content:
            thesis_md = sec.content
            highlights_raw = extract_highlights(thesis_md)

    report_headline = report_state.metadata.get("headline")
    if not report_headline:
        report_headline = f"{company_name}: Investment Analysis and Outlook"

    # ===== PAGE 1: COVER =====

    # Branded header bar (blue background, white text)
    _add_cover_header_bar(
        doc, company_name, ticker, exchange, sector, rating,
        price_target, target_upside, current_price, brand_color, hex_brand,
    )

    # ---- Two-column cover body (sidebar | main content) ----
    # Mirrors the PDF's .main-layout: sidebar 30%, main content 70%.
    cover_table = doc.add_table(rows=1, cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tbl = cover_table._tbl
    # Remove all borders for seamless layout
    _tbl_pr = _tbl.tblPr if _tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>'
    )
    _tbl_pr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    ))

    sidebar = cover_table.cell(0, 0)
    main = cover_table.cell(0, 1)
    from docx.shared import Emu
    sidebar.width = Cm(5.1)   # ~30% of 17cm usable
    main.width = Cm(11.9)     # ~70%

    # Add right border to sidebar cell for visual separation
    sb_pr = sidebar._tc.get_or_add_tcPr()
    sb_pr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tcBorders>'
    ))
    # Padding for both cells
    for c in (sidebar, main):
        c._tc.get_or_add_tcPr().append(parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            '<w:top w:w="60" w:type="dxa"/>'
            '<w:left w:w="80" w:type="dxa"/>'
            '<w:bottom w:w="60" w:type="dxa"/>'
            '<w:right w:w="120" w:type="dxa"/>'
            '</w:tcMar>'
        ))

    # ---- SIDEBAR content ----

    # 5-Year chart
    sidebar_chart_bytes = None
    try:
        from .chart_generator import generate_sidebar_chart
        sidebar_chart_bytes = generate_sidebar_chart(
            ticker, branding.firm.primary_color, period="5y"
        )
    except Exception as e:
        logger.warning("Could not generate sidebar chart for DOCX: %s", e)

    if sidebar_chart_bytes:
        sp = sidebar.paragraphs[0]
        sp.clear()
        _add_run(sp, "5-YEAR CHART", bold=True, size=Pt(8), color=brand_color)
        _set_paragraph_spacing(sp, after=Pt(4))
        # Add bottom border to mimic PDF sidebar section header
        sp._p.get_or_add_pPr().append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{hex_brand}"/>'
            '</w:pBdr>'
        ))
        cp = sidebar.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run()
        run.add_picture(io.BytesIO(sidebar_chart_bytes), width=Inches(1.8))
        _set_paragraph_spacing(cp, after=Pt(6))
    else:
        sidebar.paragraphs[0].clear()

    # Analyst info box
    ap = sidebar.add_paragraph()
    _set_cell_shading_para(ap, "F1F5F9")
    _add_run(ap, branding.analyst.full_name, bold=True, size=Pt(8), color=brand_color)
    ap2 = sidebar.add_paragraph()
    _set_cell_shading_para(ap2, "F1F5F9")
    _add_run(ap2, branding.analyst.email, size=Pt(7), color=_COLOR_SECONDARY)
    ap3 = sidebar.add_paragraph()
    _set_cell_shading_para(ap3, "F1F5F9")
    _add_run(ap3, branding.analyst.sector, size=Pt(7), color=_COLOR_SECONDARY)
    _set_paragraph_spacing(ap3, after=Pt(8))

    # Key Data section
    key_data = []
    if first_page_data:
        key_data = [
            ("Price", format_number(first_page_data.current_price, "price")),
            ("Market Cap", format_number(first_page_data.market_cap, "market_cap")),
            ("EPS", f"{first_page_data.eps:.2f}" if first_page_data.eps else "N/A"),
            ("52W High", format_number(first_page_data.week_52_high, "price")),
            ("52W Low", format_number(first_page_data.week_52_low, "price")),
        ]
        if price_target:
            key_data.append(("Price Target", f"${price_target:,.2f}"))
            if target_upside is not None:
                key_data.append(("Upside/Downside", f"{target_upside:+.1f}%"))

    if key_data:
        _add_sidebar_section(sidebar, "Key Data", key_data, brand_color, hex_brand)

    # Period Returns
    if first_page_data and first_page_data.returns:
        r = first_page_data.returns
        returns_data = []
        for period, val in [("1M", r.return_1m), ("YTD", r.return_ytd),
                            ("1Y", r.return_1y), ("5Y", r.return_5y), ("10Y", r.return_10y)]:
            if val is not None:
                returns_data.append((period, f"{val:+.1f}%"))
        if returns_data:
            _add_sidebar_section(sidebar, "Period Returns", returns_data, brand_color, hex_brand)

    # Valuation ratios
    if first_page_data:
        val_data = [
            ("P/E (TTM)", format_number(first_page_data.pe_ratio, "ratio")),
            ("EV/EBITDA", format_number(first_page_data.ev_to_ebitda, "ratio")),
            ("P/B", format_number(first_page_data.price_to_book, "ratio")),
            ("Beta", f"{first_page_data.beta:.2f}" if first_page_data.beta else "N/A"),
            ("Div Yield", format_number(first_page_data.dividend_yield, "percent_raw")),
            ("ROE", format_number(first_page_data.roe, "percent_raw")),
        ]
        _add_sidebar_section(sidebar, "Valuation", val_data, brand_color, hex_brand)

    # ---- MAIN CONTENT ----

    # Headline
    mp = main.paragraphs[0]
    mp.clear()
    _add_run(mp, report_headline, bold=True, size=Pt(14), color=brand_color)
    _set_paragraph_spacing(mp, after=Pt(10))

    # Key highlights in accent box
    if highlights_raw:
        hp = main.add_paragraph()
        _add_run(hp, "KEY HIGHLIGHTS", bold=True, size=Pt(9), color=brand_color)
        _set_paragraph_spacing(hp, after=Pt(4))
        # Left border accent on each highlight
        for h in highlights_raw[:5]:
            bp = main.add_paragraph()
            _add_run(bp, "\u2022  ", size=_FONT_SIZE_BODY, color=_COLOR_SECONDARY)
            _write_inline_formatting(bp, h)
            _set_paragraph_spacing(bp, before=Pt(1), after=Pt(1))

    # Investment thesis summary on cover (short preview, full version in body)
    if thesis_md:
        main.add_paragraph()  # spacing
        tp = main.add_paragraph()
        _add_run(tp, "Investment Summary", bold=True, size=_FONT_SIZE_H2, color=brand_color)
        _set_paragraph_spacing(tp, after=Pt(6))
        intro_text = thesis_md[:600]
        boundary = intro_text.rfind("\n\n", 300)
        if boundary > 0:
            intro_text = intro_text[:boundary]
        _write_markdown_to_doc(main, intro_text, brand_color)

    # Page break after cover
    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    doc.add_paragraph("Contents", style='Heading 1')
    toc_para = doc.add_paragraph()
    _add_run(toc_para, "Right-click and select 'Update Field' to generate table of contents.",
             size=_FONT_SIZE_SMALL, color=_COLOR_SECONDARY, italic=True)

    # Insert TOC field code (same pattern as PAGE field in _add_page_number_run)
    from docx.oxml import OxmlElement
    toc_run = toc_para.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'begin')
    toc_run._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    toc_run._r.append(instr)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'end')
    toc_run._r.append(fld_end)

    doc.add_page_break()

    # ===== BODY: Prose sections =====

    section_order = [
        ("investment_thesis", "Investment Thesis"),
        ("fundamentals", "Fundamental Analysis"),
        ("moat", "Competitive Moat"),
        ("strategy", "Strategic Assessment"),
        ("industry", "Industry Dynamics"),
    ]

    first_body_section = True
    for key, title in section_order:
        section = report_state.sections.get(key)
        if section and hasattr(section, "content") and section.content:
            _add_section_heading(doc, title, brand_color,
                                 page_break=(not first_body_section))
            _write_markdown_to_doc(doc, section.content, brand_color)
            first_body_section = False

    # Bull / Bear cases
    bull_sec = report_state.sections.get("bull_case")
    bear_sec = report_state.sections.get("bear_case")
    if (bull_sec and hasattr(bull_sec, "content") and bull_sec.content) or \
       (bear_sec and hasattr(bear_sec, "content") and bear_sec.content):
        _add_section_heading(doc, "Bull & Bear Cases", brand_color,
                             page_break=(not first_body_section))

        if bull_sec and hasattr(bull_sec, "content") and bull_sec.content:
            cell = _add_accent_box(doc, bg_hex="#dcfce7", border_hex="#16a34a")
            p = cell.paragraphs[0]
            _add_run(p, "Bull Case", bold=True, size=_FONT_SIZE_H3, color=_COLOR_POSITIVE)
            _set_paragraph_spacing(p, before=Pt(4), after=Pt(4))
            _write_markdown_to_doc(cell, bull_sec.content, brand_color)

        if bear_sec and hasattr(bear_sec, "content") and bear_sec.content:
            cell = _add_accent_box(doc, bg_hex="#fee2e2", border_hex="#dc2626")
            p = cell.paragraphs[0]
            _add_run(p, "Bear Case", bold=True, size=_FONT_SIZE_H3, color=_COLOR_NEGATIVE)
            _set_paragraph_spacing(p, before=Pt(4), after=Pt(4))
            _write_markdown_to_doc(cell, bear_sec.content, brand_color)

    # ===== VALUATION SECTIONS =====

    # Football Field / Valuation Summary
    football_ranges = val_display.football_ranges
    if football_ranges:
        _add_section_heading(doc, "Valuation Summary", brand_color, page_break=True)
        _add_football_field_table(doc, football_ranges,
                                  val_display.football_current_price, brand_color)

    # DCF
    dcf_summary = val_display.dcf_summary
    if dcf_summary:
        _add_section_heading(doc, "DCF Valuation", brand_color, page_break=True)
        dcf_sec = report_state.sections.get("dcf")
        if dcf_sec and hasattr(dcf_sec, "content") and dcf_sec.content:
            _write_markdown_to_doc(doc, strip_markdown_tables(dcf_sec.content), brand_color)
        _add_dcf_summary_table(doc, dcf_summary, brand_color)

    # Scenario Analysis
    scenario_summary = val_display.scenario_summary
    if scenario_summary:
        _add_section_heading(doc, "Scenario Analysis", brand_color, page_break=True)
        scen_sec = report_state.sections.get("scenarios")
        if scen_sec and hasattr(scen_sec, "content") and scen_sec.content:
            _write_markdown_to_doc(doc, strip_markdown_tables(scen_sec.content), brand_color)
        _add_scenario_table(doc, scenario_summary, brand_color)

    # Comparable Companies
    comp_summary = val_display.comp_summary
    if comp_summary:
        _add_section_heading(doc, "Comparable Companies", brand_color, page_break=True)
        comps_sec = report_state.sections.get("comps")
        if comps_sec and hasattr(comps_sec, "content") and comps_sec.content:
            _write_markdown_to_doc(doc, strip_markdown_tables(comps_sec.content), brand_color)
        _add_comps_table(doc, comp_summary, brand_color)

    # Precedent Transactions
    precedent_summary = val_display.precedent_summary
    if precedent_summary:
        _add_section_heading(doc, "Precedent Transactions", brand_color)
        prec_sec = report_state.sections.get("precedents")
        if prec_sec and hasattr(prec_sec, "content") and prec_sec.content:
            _write_markdown_to_doc(doc, strip_markdown_tables(prec_sec.content), brand_color)

        deals = precedent_summary.get("deals", [])
        if deals:
            headers = ["Target", "Acquirer", "Date", "EV/Revenue", "EV/EBITDA", "Premium"]
            rows = []
            for d in deals:
                rows.append([
                    d.get("target", ""),
                    d.get("acquirer", ""),
                    d.get("date", ""),
                    f"{d['ev_to_revenue']:.1f}x" if d.get("ev_to_revenue") else "N/A",
                    f"{d['ev_to_ebitda']:.1f}x" if d.get("ev_to_ebitda") else "N/A",
                    f"{d['premium']:.1f}%" if d.get("premium") else "N/A",
                ])
            _add_data_table(doc, headers, rows, brand_color)

    # Earnings Model — prose and structured data rendered independently
    earnings_table = val_display.earnings_table
    earn_sec = report_state.sections.get("earnings_model")
    has_earnings_prose = earn_sec and hasattr(earn_sec, "content") and earn_sec.content
    if has_earnings_prose or earnings_table:
        _add_section_heading(doc, "Earnings Model", brand_color)
        if has_earnings_prose:
            _write_markdown_to_doc(doc, strip_markdown_tables(earn_sec.content), brand_color)
        if earnings_table:
            _add_earnings_table(doc, earnings_table, brand_color)

    # Sensitivity Analysis
    sensitivity_grid = val_display.sensitivity_grid
    if sensitivity_grid:
        _add_section_heading(doc, "Sensitivity Analysis", brand_color)
        _add_sensitivity_grid_table(doc, sensitivity_grid, brand_color,
                                    title="WACC vs Terminal Growth Rate")

    growth_margin_grid = val_display.growth_margin_grid
    if growth_margin_grid:
        _add_sensitivity_grid_table(doc, growth_margin_grid, brand_color,
                                    title="Revenue Growth vs Operating Margin")

    # Conviction & Rating
    conviction_data = val_display.conviction_structured
    conviction_categories = val_display.conviction_categories
    if conviction_data:
        _add_section_heading(doc, "Conviction & Rating", brand_color)
        conv_sec = report_state.sections.get("conviction")
        if conv_sec and hasattr(conv_sec, "content") and conv_sec.content:
            _write_markdown_to_doc(doc, strip_markdown_tables(conv_sec.content), brand_color)

        if conviction_categories:
            headers = ["Category", "Score", "Evidence"]
            rows = []
            for cat in conviction_categories:
                rows.append([
                    cat.get("name", ""),
                    f"{cat['score']:.1f}/100" if cat.get("score") is not None else "N/A",
                    cat.get("evidence", ""),
                ])
            _add_data_table(doc, headers, rows, brand_color)

    # Financial Statements
    if financial_statements:
        _add_section_heading(doc, "Financial Statements", brand_color, page_break=True)
        _add_financial_statements(doc, financial_statements, brand_color)

    # Technical Analysis
    tech_sec = report_state.sections.get("technicals")
    if tech_sec and hasattr(tech_sec, "content") and tech_sec.content:
        _add_section_heading(doc, "Technical Analysis", brand_color, page_break=True)
        _write_markdown_to_doc(doc, tech_sec.content, brand_color)

        # Technical chart
        try:
            from .chart_generator import generate_technical_chart
            chart_bytes = generate_technical_chart(ticker, branding.firm.primary_color, period="1y")
            if chart_bytes:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(io.BytesIO(chart_bytes), width=Inches(5.5))
        except Exception as e:
            logger.warning("Could not generate technical chart for DOCX: %s", e)

    # ===== SOURCES & REFERENCES =====

    sources_section = report_state.sections.get("sources")
    if sources_section and hasattr(sources_section, "sources") and sources_section.sources:
        _add_section_heading(doc, "Sources & References", brand_color, page_break=True)
        for source in sources_section.sources:
            if isinstance(source, Source):
                p = doc.add_paragraph()
                _add_run(p, f"[{source.id}] ", bold=True, size=_FONT_SIZE_SMALL, color=_COLOR_BODY)
                _add_run(p, source.title, size=_FONT_SIZE_SMALL, color=_COLOR_BODY)
                if source.url:
                    _add_run(p, f" — {source.url}", size=_FONT_SIZE_SMALL, color=_COLOR_SECONDARY)
                if source.date and source.date != "N/A":
                    _add_run(p, f" ({source.date})", size=_FONT_SIZE_SMALL, color=_COLOR_SECONDARY)
                _set_paragraph_spacing(p, after=Pt(2))
            elif isinstance(source, dict):
                p = doc.add_paragraph()
                _add_run(p, f"[{source.get('id', '')}] ", bold=True, size=_FONT_SIZE_SMALL, color=_COLOR_BODY)
                _add_run(p, source.get("title", ""), size=_FONT_SIZE_SMALL, color=_COLOR_BODY)
                _set_paragraph_spacing(p, after=Pt(2))

    # ===== DISCLAIMER =====

    doc.add_page_break()
    _add_section_heading(doc, "Important Disclosures", brand_color)

    disclaimer_text = (
        f"This report was prepared by {branding.firm.name} and is intended for informational "
        f"purposes only. It does not constitute investment advice, a recommendation, or a "
        f"solicitation to buy or sell any security. The information contained herein is based "
        f"on sources believed to be reliable, but no representation or warranty is made as to "
        f"its accuracy or completeness. Past performance is not indicative of future results. "
        f"Investors should conduct their own due diligence before making any investment decision."
    )
    p = doc.add_paragraph()
    _add_run(p, disclaimer_text, size=_FONT_SIZE_SMALL, color=_COLOR_SECONDARY)
    _set_paragraph_spacing(p, after=Pt(12))

    # Analyst attribution
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, branding.analyst.full_name, bold=True, size=Pt(12), color=brand_color)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p2, branding.analyst.email, size=_FONT_SIZE_SMALL, color=_COLOR_SECONDARY)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p3, before=Pt(20))
    _add_run(p3, branding.firm.name, bold=True, size=Pt(14), color=brand_color)

    # Generation metadata
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, before=Pt(20))
    _add_run(p, f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} | {branding.firm.tool_branding}",
             size=Pt(7), color=_COLOR_SECONDARY)

    # ===== Save to bytes =====

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
